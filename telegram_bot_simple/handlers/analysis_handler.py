"""
Обработчик анализа CSV файлов с Claude AI
"""

import os
import sys
import logging
import asyncio
from pathlib import Path
from datetime import datetime
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import ContextTypes
from telegram.constants import ParseMode

# Импорты для анализа
import csv
import pandas as pd
from anthropic import Anthropic
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)


class SimpleClaudeAnalyzer:
    """Простой анализатор с Claude AI без зависимостей от services"""

    def __init__(self, api_key: str):
        self.client = Anthropic(api_key=api_key)
        self.model = "claude-sonnet-4-20250514"

    def analyze_and_generate_report(self, csv_file: str, output_file: str) -> str:
        """
        Анализ CSV файла и генерация DOCX отчета

        Args:
            csv_file: Путь к CSV файлу
            output_file: Путь для сохранения DOCX

        Returns:
            str: Путь к созданному DOCX файлу
        """
        # Читаем CSV
        df = pd.read_csv(csv_file, encoding='utf-8')

        # Подготавливаем данные для анализа
        messages_text = "\n".join([
            f"[{row['Date']}] {row['From']}: {row['Text']}"
            for _, row in df.head(3000).iterrows()
        ])

        # Создаем промпт
        prompt = f"""Проанализируй следующие сообщения из Telegram чата и создай детальный отчет.

СООБЩЕНИЯ:
{messages_text}

Создай структурированный анализ, включающий:
1. Общая статистика и обзор
2. Ключевые темы обсуждений
3. Основные участники и их роли
4. Тренды и паттерны
5. Важные инсайты и выводы

Формат ответа: структурированный текст с заголовками."""

        # Отправляем запрос к Claude
        response = self.client.messages.create(
            model=self.model,
            max_tokens=4096,
            messages=[{
                "role": "user",
                "content": prompt
            }]
        )

        analysis_text = response.content[0].text

        # Создаем DOCX отчет
        doc = Document()

        # Заголовок
        title = doc.add_heading('Анализ Telegram переписки', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Дата создания
        date_para = doc.add_paragraph(f'Дата создания: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()  # Пустая строка

        # Статистика
        stats_heading = doc.add_heading('📊 Статистика', 1)
        doc.add_paragraph(f'Всего сообщений: {len(df):,}')
        doc.add_paragraph(f'Уникальных авторов: {df["From"].nunique():,}')
        doc.add_paragraph(f'Период: {df["Date"].min()} - {df["Date"].max()}')

        doc.add_paragraph()

        # Анализ от Claude
        analysis_heading = doc.add_heading('🤖 Анализ Claude AI', 1)

        # Разбиваем анализ на параграфы
        for line in analysis_text.split('\n'):
            if line.strip():
                if line.startswith('#'):
                    # Это заголовок
                    level = line.count('#')
                    doc.add_heading(line.replace('#', '').strip(), level)
                else:
                    doc.add_paragraph(line.strip())

        # Сохраняем
        doc.save(output_file)
        return output_file


class AnalysisHandler:
    """Обработчик анализа с Claude AI"""

    def __init__(self, user_manager):
        self.user_manager = user_manager

    async def start_analysis(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        """Начать процесс анализа"""
        query = update.callback_query
        await query.answer()

        user_id = update.effective_user.id

        # Проверка настроек
        if not self.user_manager.has_claude_config(user_id):
            keyboard = [[InlineKeyboardButton("⚙️ Настроить", callback_data="settings")]]
            reply_markup = InlineKeyboardMarkup(keyboard)

            await query.edit_message_text(
                "⚠️ **Claude API не настроен**\n\n"
                "Сначала настройте Claude API ключ в разделе Настройки.",
                reply_markup=reply_markup,
                parse_mode=ParseMode.MARKDOWN
            )
            return

        keyboard = [[InlineKeyboardButton("« Отмена", callback_data="main_menu")]]
        reply_markup = InlineKeyboardMarkup(keyboard)

        message = """
🤖 **Анализ с Claude AI**

**Как использовать:**
1. Отправьте CSV файл с сообщениями
2. Дождитесь анализа
3. Получите DOCX отчет

**Требования к файлу:**
• Формат: CSV
• Кодировка: UTF-8
• Столбцы: Date, From, Text
• Максимум: 3000 строк

**Что делает анализ:**
• Определяет ключевые темы
• Выявляет тренды обсуждений
• Анализирует тональность
• Создает профессиональный отчет

Отправьте CSV файл для анализа:
"""

        await query.edit_message_text(
            message,
            reply_markup=reply_markup,
            parse_mode=ParseMode.MARKDOWN
        )

        # Сохраняем состояние
        context.user_data['state'] = 'analysis_file_upload'

    async def handle_analysis_callback(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        """Обработка callback кнопок анализа"""
        # Пока не используется
        pass

    async def handle_text_input(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        """Обработка текстового ввода (если нужно)"""
        state = context.user_data.get('state')

        if state == 'analysis_file_upload':
            await update.message.reply_text(
                "⚠️ Пожалуйста, отправьте CSV файл (не текст).\n\n"
                "Прикрепите файл как документ."
            )

    async def handle_document(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        """Обработка загруженного документа"""
        user_id = update.effective_user.id
        state = context.user_data.get('state')

        # Проверяем состояние
        if state != 'analysis_file_upload':
            return

        document = update.message.document

        # Проверяем что это CSV файл
        if not document.file_name.endswith('.csv'):
            await update.message.reply_text(
                "❌ Это не CSV файл!\n\n"
                "Отправьте файл с расширением .csv"
            )
            return

        # Проверяем размер (макс 20MB)
        if document.file_size > 20 * 1024 * 1024:
            await update.message.reply_text(
                "❌ Файл слишком большой!\n\n"
                "Максимальный размер: 20 MB"
            )
            return

        # Очищаем состояние
        context.user_data.pop('state', None)

        # Отправляем сообщение о начале анализа
        status_message = await update.message.reply_text(
            "📥 **Загружаю файл...**\n\n"
            f"Файл: `{document.file_name}`\n"
            f"Размер: {document.file_size / 1024:.1f} KB",
            parse_mode=ParseMode.MARKDOWN
        )

        try:
            # Создаем директорию для загрузок
            upload_dir = Path(__file__).parent.parent / "data" / "uploads" / str(user_id)
            upload_dir.mkdir(parents=True, exist_ok=True)

            # Скачиваем файл
            file = await document.get_file()
            file_path = upload_dir / document.file_name
            await file.download_to_drive(str(file_path))

            # Запускаем анализ
            await self._run_analysis(update, context, status_message, file_path)

        except Exception as e:
            logger.error(f"Error handling document: {e}")
            await status_message.edit_text(
                f"❌ **Ошибка загрузки файла**\n\n"
                f"Причина: {str(e)}",
                parse_mode=ParseMode.MARKDOWN
            )

    async def _run_analysis(self, update: Update, context: ContextTypes.DEFAULT_TYPE,
                           status_message, csv_file: Path):
        """Выполнение анализа"""
        user_id = update.effective_user.id

        try:
            # Получаем настройки
            settings = self.user_manager.get_settings(user_id)
            claude_api_key = settings['claude_api_key']

            # Проверяем количество строк
            await status_message.edit_text(
                "🔍 **Проверяю файл...**",
                parse_mode=ParseMode.MARKDOWN
            )

            import csv
            with open(csv_file, 'r', encoding='utf-8') as f:
                reader = csv.reader(f)
                header = next(reader)
                row_count = sum(1 for _ in reader)

            if row_count > 3000:
                await status_message.edit_text(
                    f"⚠️ **Файл слишком большой**\n\n"
                    f"Строк в файле: {row_count:,}\n"
                    f"Максимум: 3,000\n\n"
                    f"Пожалуйста, уменьшите файл (используйте фильтр по датам при экспорте).",
                    parse_mode=ParseMode.MARKDOWN
                )
                return

            # Обновляем статус
            await status_message.edit_text(
                "🤖 **Анализирую с Claude AI...**\n\n"
                f"Строк: {row_count:,}\n\n"
                "⏳ Это может занять 1-2 минуты...\n"
                "Пожалуйста, подождите.",
                parse_mode=ParseMode.MARKDOWN
            )

            # Создаем анализатор
            analyzer = SimpleClaudeAnalyzer(api_key=claude_api_key)

            # Генерируем имя выходного файла
            output_dir = Path(__file__).parent.parent / "data" / "reports" / str(user_id)
            output_dir.mkdir(parents=True, exist_ok=True)

            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_file = output_dir / f"analysis_{timestamp}.docx"

            # Запускаем анализ
            # Создаем wrapper для синхронных операций в асинхронном контексте
            loop = asyncio.get_event_loop()
            result_file = await loop.run_in_executor(
                None,
                analyzer.analyze_and_generate_report,
                str(csv_file),
                str(output_file)
            )

            # Проверяем результат
            if not os.path.exists(result_file):
                raise Exception("Файл отчета не был создан")

            # Получаем размер файла
            file_size = os.path.getsize(result_file)
            file_size_kb = file_size / 1024

            # Отправляем файл пользователю
            await status_message.edit_text(
                "📤 **Отправляю отчет...**",
                parse_mode=ParseMode.MARKDOWN
            )

            with open(result_file, 'rb') as f:
                await update.message.reply_document(
                    document=f,
                    filename=os.path.basename(result_file),
                    caption=f"✅ **Анализ завершен!**\n\n"
                            f"📊 Статистика:\n"
                            f"• Проанализировано строк: {row_count:,}\n"
                            f"• Размер отчета: {file_size_kb:.1f} KB\n"
                            f"• Модель: Claude Sonnet 4",
                    parse_mode=ParseMode.MARKDOWN
                )

            # Удаляем статусное сообщение
            await status_message.delete()

            # Показываем меню
            keyboard = [[InlineKeyboardButton("« Главное меню", callback_data="main_menu")]]
            reply_markup = InlineKeyboardMarkup(keyboard)

            await update.message.reply_text(
                "🎉 Анализ успешно завершен!\n\n"
                "Откройте DOCX файл для просмотра детального отчета.",
                reply_markup=reply_markup
            )

        except Exception as e:
            logger.error(f"Analysis error: {e}")

            error_message = str(e)
            if "API key" in error_message or "authentication" in error_message.lower():
                error_message = "Ошибка API ключа Claude. Проверьте настройки."
            elif "rate limit" in error_message.lower():
                error_message = "Превышен лимит запросов к API. Попробуйте позже."
            elif "timeout" in error_message.lower():
                error_message = "Превышено время ожидания. Попробуйте файл меньшего размера."

            await status_message.edit_text(
                f"❌ **Ошибка анализа**\n\n"
                f"Причина: {error_message}\n\n"
                "Попробуйте:\n"
                "1. Проверить формат CSV файла\n"
                "2. Уменьшить размер файла\n"
                "3. Проверить настройки Claude API",
                parse_mode=ParseMode.MARKDOWN
            )
