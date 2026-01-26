# services/analyzer.py
"""Модуль для анализа CSV файлов с помощью Claude API (Anthropic)"""

import os
import time
import pandas as pd
import logging
from pathlib import Path
from datetime import datetime
from typing import Optional

import anthropic
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from core.config import (
    CLAUDE_API_KEY,
    get_input_folder,
    get_output_folder,
    get_logs_dir
)

# Настройка логирования
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Константы Claude API
CLAUDE_MODEL = "claude-sonnet-4-20250514"  # Актуальная модель
MAX_TOKENS = 8192  # Максимум токенов в ответе
MAX_RETRIES = 3  # Количество попыток при ошибке API
MAX_CSV_ROWS = 3000  # Лимит строк CSV для контекста
API_DELAY_SECONDS = 5  # Задержка между запросами к API (секунды)

def get_client(api_key: Optional[str] = None) -> anthropic.Anthropic:
    """
    Получение или создание клиента Claude API

    Args:
        api_key: Claude API ключ (если None, использует глобальный из конфига)

    Returns:
        Клиент Claude API
    """
    # Если передан api_key, создаем новый клиент для этого ключа
    if api_key and api_key.strip():
        return anthropic.Anthropic(api_key=api_key)

    # Иначе используем глобальный ключ из конфига
    if not CLAUDE_API_KEY or CLAUDE_API_KEY.strip() == "":
        raise ValueError(
            "CLAUDE_API_KEY не настроен. Откройте настройки и введите ключ API."
        )

    logger.info("✅ Using global Claude API key from config")
    return anthropic.Anthropic(api_key=CLAUDE_API_KEY)


def analyze_csv_with_claude(file_path: str, claude_api_key: Optional[str] = None, custom_prompt: Optional[str] = None) -> str:
    """
    Читает CSV и отправляет данные в Claude API для анализа.

    Args:
        file_path: Путь к CSV файлу
        claude_api_key: Claude API ключ (опционально, если None - использует глобальный)
        custom_prompt: Кастомный промпт пользователя (опционально, если None - использует дефолтный)

    Returns:
        Текст анализа от Claude

    Raises:
        ValueError: Если Claude API ключ не настроен
        Exception: Другие ошибки при анализе
    """
    try:
        # Создать клиент с переданным или глобальным ключом
        client = get_client(api_key=claude_api_key)
        logger.info(f"📖 Reading file: {file_path}")

        # Чтение CSV с поддержкой разных форматов
        df = _read_csv_flexible(file_path)
        logger.info(f"✅ Файл прочитан. Строк: {len(df)}, Столбцов: {len(df.columns)}")

        # Проверка обязательных колонок
        required_columns = ['Date', 'From', 'Text']
        missing = [col for col in required_columns if col not in df.columns]
        if missing:
            return f"Ошибка: Отсутствуют обязательные колонки: {missing}"

        # Ограничение размера для API
        if len(df) > MAX_CSV_ROWS:
            logger.warning(
                f"Файл содержит {len(df)} строк. "
                f"Ограничиваем до {MAX_CSV_ROWS} для анализа."
            )
            df = df.head(MAX_CSV_ROWS)

        # Формирование контента
        csv_content = df.to_string(index=False)

        # Промпт для анализа (используем кастомный если есть, иначе дефолтный)
        if custom_prompt:
            logger.info("🎯 Используется кастомный промпт пользователя")
            # В кастомном промпте {csv_content} будет заменен на данные
            prompt = custom_prompt.replace("{csv_content}", csv_content)
        else:
            logger.info("📝 Используется дефолтный промпт анализа")
            prompt = _build_analysis_prompt(csv_content)

        logger.info("🤖 Отправка запроса в Claude API...")

        # Отправка с повторами при ошибке
        for attempt in range(MAX_RETRIES):
            try:
                message = client.messages.create(
                    model=CLAUDE_MODEL,
                    max_tokens=MAX_TOKENS,
                    messages=[
                        {"role": "user", "content": prompt}
                    ]
                )

                logger.info("✅ Ответ получен от Claude API")

                # Извлекаем текст из ответа
                if message.content and len(message.content) > 0:
                    return message.content[0].text
                else:
                    return "Ошибка: Пустой ответ от Claude API"

            except anthropic.RateLimitError as e:
                wait_time = 2 ** (attempt + 1)
                logger.warning(
                    f"Rate limit. Попытка {attempt + 1}/{MAX_RETRIES}. "
                    f"Ожидание {wait_time} сек..."
                )
                time.sleep(wait_time)

            except anthropic.APIError as e:
                logger.warning(f"API ошибка: {e}. Попытка {attempt + 1}/{MAX_RETRIES}")
                if attempt == MAX_RETRIES - 1:
                    raise
                time.sleep(2 ** attempt)

        return "Ошибка: Превышено количество попыток обращения к API"

    except Exception as e:
        error_msg = f"Ошибка при анализе файла {file_path}: {e}"
        logger.error(error_msg)
        return error_msg


def _read_csv_flexible(file_path: str) -> pd.DataFrame:
    """Чтение CSV с автоопределением формата"""
    encodings = ['utf-8-sig', 'utf-8', 'cp1251', 'latin-1']
    separators = [';', ',', '\t']

    for encoding in encodings:
        for sep in separators:
            try:
                df = pd.read_csv(file_path, sep=sep, encoding=encoding)
                # Проверяем, что получили больше одной колонки
                if len(df.columns) > 1:
                    return df
            except Exception:
                continue

    # Последняя попытка с автоопределением
    return pd.read_csv(file_path, sep=None, encoding='utf-8-sig', engine='python')


def _build_analysis_prompt(csv_content: str) -> str:
    """Построение промпта для анализа"""
    return f"""Проанализируй следующие данные из CSV файла, содержащие информацию о взаимодействиях с клиентами и работе менеджеров службы поддержки.

Менеджеры в CSV файлах имеют наименование, которое содержит в себе 'Fulfillment-Box Support', 'Support', 'Fulfillment-Box' и подобное.

Данные CSV файла:
{csv_content}

На основе этих данных предоставь анализ по следующим пунктам:

1. **Самые частые запросы клиентов** — повторяющиеся вопросы и темы обращений, что позволит доработать Ysell и другие процессы.

2. **Причины конфликтов и недовольства клиентов** — выяви основные проблемные области.

3. **Количество вопросов по каждому менеджеру** — статистика обработанных обращений.

4. **Среднее время ответа менеджеров** — отдельно для будней и выходных.

Требования к анализу времени ответа:
- Раздели сообщения по календарным дням с указанием дня недели
- Для каждого клиентского сообщения найди первое последующее сообщение менеджера в течение того же дня
- Исключи случаи, когда первое последующее сообщение — это системная/сервисная отправка не от support-аккаунта
- Рассчитай среднее и медиану времени ответа отдельно для будней и выходных

Формат вывода:
- Структурированный текст без таблиц (для удобного просмотра в DOCX)
- Без упоминания номеров строк и сообщений
- Не включай методику обработки в отчёт
- Не добавляй рекомендации и советы в конце
- Это только аналитический отчёт

Предоставь анализ на русском языке."""


def save_to_docx(text_content: str, output_file_path: str, source_filename: str):
    """
    Сохраняет текстовый контент в DOCX файл с форматированием.

    Args:
        text_content: Текст для сохранения
        output_file_path: Путь к выходному файлу
        source_filename: Имя исходного файла для заголовка
    """
    try:
        logger.info(f"💾 Сохранение результатов в {output_file_path}")

        doc = Document()

        # Заголовок
        title = doc.add_heading(f'Анализ файла: {source_filename}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Дата создания
        doc.add_paragraph(
            f"Дата создания отчёта: {datetime.now().strftime('%d.%m.%Y %H:%M:%S')}"
        )
        doc.add_paragraph(f"Модель анализа: {CLAUDE_MODEL}")
        doc.add_paragraph()

        # Обработка текста
        for line in text_content.split('\n'):
            line = line.strip()
            if not line:
                continue

            # Определение типа строки
            if line.startswith('**') and line.endswith('**'):
                # Жирный текст как заголовок
                doc.add_heading(line.strip('*').strip(), level=2)
            elif line[0].isdigit() and '.' in line[:3] and '**' in line:
                # Нумерованный заголовок с жирным текстом
                doc.add_heading(line.replace('**', '').strip(), level=2)
            elif line.startswith('###'):
                doc.add_heading(line.replace('#', '').strip(), level=3)
            elif line.startswith('##'):
                doc.add_heading(line.replace('#', '').strip(), level=2)
            elif line.startswith('#'):
                doc.add_heading(line.replace('#', '').strip(), level=1)
            else:
                paragraph = doc.add_paragraph()

                # Обработка жирного текста внутри строки
                if '**' in line:
                    parts = line.split('**')
                    for i, part in enumerate(parts):
                        if part:
                            run = paragraph.add_run(part)
                            if i % 2 == 1:  # Нечётные части — жирные
                                run.bold = True
                else:
                    paragraph.add_run(line)

                # Отступ для списков
                if line.startswith('-') or line.startswith('•') or line.startswith('*'):
                    paragraph.paragraph_format.left_indent = Pt(20)

        doc.save(output_file_path)
        logger.info(f"✅ Файл сохранён: {output_file_path}")

    except Exception as e:
        logger.error(f"❌ Ошибка при сохранении DOCX: {e}")
        raise


def analyze_csv_folder(
        input_folder: Optional[str] = None,
        output_folder: Optional[str] = None,
        progress_callback: Optional[callable] = None
) -> dict:
    """
    Анализирует все CSV файлы в папке и создаёт DOCX отчёты.

    Args:
        input_folder: Путь к папке с CSV (по умолчанию из конфига)
        output_folder: Путь к папке для DOCX (по умолчанию из конфига)
        progress_callback: Функция обратного вызова для обновления прогресса
                          Сигнатура: callback(current: int, total: int, filename: str, status: str)

    Returns:
        Словарь с результатами: {'success': int, 'errors': int, 'details': list}
    """
    # Используем пути из конфига если не указаны
    if input_folder is None:
        input_folder = str(get_input_folder())
    if output_folder is None:
        output_folder = str(get_output_folder())

    logger.info("\n" + "=" * 50)
    logger.info("🔍 НАЧАЛО АНАЛИЗА CSV ФАЙЛОВ (Claude API)")
    logger.info("=" * 50 + "\n")

    # Проверка API ключа
    if not CLAUDE_API_KEY or CLAUDE_API_KEY.strip() == "":
        error_msg = "❌ CLAUDE_API_KEY не настроен. Откройте настройки и введите ключ API."
        logger.error(error_msg)
        raise ValueError(error_msg)

    logger.info(f"📁 Входная папка: {input_folder}")
    logger.info(f"📁 Выходная папка: {output_folder}")
    logger.info(f"🤖 Модель: {CLAUDE_MODEL}\n")

    # Проверка входной папки
    input_path = Path(input_folder)
    if not input_path.exists():
        error_msg = f"❌ Входная папка не найдена: {input_folder}"
        logger.error(error_msg)
        raise FileNotFoundError(error_msg)

    # Создание выходной папки
    output_path = Path(output_folder)
    output_path.mkdir(parents=True, exist_ok=True)

    # Получение списка CSV файлов
    # СТАЛО (без дубликатов):
    csv_files_lower = list(input_path.glob("*.csv"))
    csv_files_upper = list(input_path.glob("*.CSV"))
    all_csv_files = {f.resolve() for f in csv_files_lower + csv_files_upper}  # set убирает дубли
    csv_files = sorted(all_csv_files, key=lambda f: f.name)

    if not csv_files:
        logger.warning(f"⚠️ В папке {input_folder} не найдено CSV файлов")
        return {'success': 0, 'errors': 0, 'details': []}

    logger.info(f"📊 Найдено файлов для анализа: {len(csv_files)}\n")

    # Обработка файлов
    success_count = 0
    error_count = 0
    error_details = []

    for idx, csv_file in enumerate(csv_files, 1):
        filename = csv_file.name

        logger.info(f"\n{'=' * 50}")
        logger.info(f"📄 [{idx}/{len(csv_files)}] Обработка: {filename}")
        logger.info("=" * 50)

        # Колбэк прогресса: начало обработки файла
        if progress_callback:
            progress_callback(idx, len(csv_files), filename, "analyzing")

        try:
            # Анализ файла
            analysis_result = analyze_csv_with_claude(str(csv_file))

            # Проверка на ошибку
            if analysis_result.startswith("Ошибка"):
                logger.warning(f"⚠️ Ошибка анализа {filename}")
                error_count += 1
                error_details.append(f"{filename}: {analysis_result}")
                continue

            # Формирование выходного файла
            output_filename = csv_file.stem + "_analysis.docx"
            output_file_path = output_path / output_filename

            # Сохранение результата
            save_to_docx(analysis_result, str(output_file_path), filename)

            # Удаление исходного файла после успешного сохранения
            try:
                csv_file.unlink()
                logger.info(f"🗑️ Исходный файл {filename} удалён")
            except Exception as e:
                logger.warning(f"⚠️ Не удалось удалить {filename}: {e}")

            success_count += 1
            logger.info(f"✅ Файл {filename} успешно обработан!")

            # Задержка перед следующим файлом (для rate limit)
            if idx < len(csv_files):
                logger.info(f"⏳ Пауза {API_DELAY_SECONDS} сек. перед следующим файлом...")

                # Колбэк прогресса: ожидание
                if progress_callback:
                    progress_callback(idx, len(csv_files), filename, "waiting")

                time.sleep(API_DELAY_SECONDS)

        except Exception as e:
            error_msg = f"❌ Ошибка при обработке {filename}: {e}"
            logger.error(error_msg)
            error_count += 1
            error_details.append(f"{filename}: {str(e)}")

    # Итоговая статистика
    logger.info("\n" + "=" * 50)
    logger.info("📊 РЕЗУЛЬТАТЫ АНАЛИЗА")
    logger.info("=" * 50)
    logger.info(f"✅ Успешно обработано: {success_count}")
    logger.info(f"❌ Ошибок: {error_count}")

    if error_details:
        logger.info("\nДетали ошибок:")
        for detail in error_details:
            logger.info(f"  - {detail}")

    logger.info(f"\n📁 Результаты сохранены в: {output_folder}")
    logger.info("=" * 50 + "\n")

    if success_count == 0 and csv_files:
        raise Exception("Не удалось обработать ни одного файла")

    return {
        'success': success_count,
        'errors': error_count,
        'details': error_details
    }